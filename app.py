from docx import Document

doc = Document()
doc.add_heading("My First Word File", level=1)
doc.add_paragraph("This line was typed by Python.")
doc.save("my_file.docx")
